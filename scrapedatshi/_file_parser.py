"""
scrapedatshi._file_parser
~~~~~~~~~~~~~~~~~~~~~~~~~
Local file text extraction helpers.

These functions run on the CLIENT machine — not on the scrapedatshi server.
They extract plain text from local files so the heavy CPU work stays off
the server. The extracted text is then submitted to /v1/process-text for
chunking.

Supported formats:
    .pdf              — pdfplumber (text layer) with basic fallback
    .md, .txt, .text  — read as-is
    .yaml, .yml       — YAML → formatted text
    .json             — JSON → formatted text
    .csv              — CSV rows → text blocks
    .xlsx, .xls       — Excel sheets → text (requires openpyxl)
    .docx             — Word documents (requires python-docx)
    .ipynb            — Jupyter notebooks (code + markdown cells)
    .html, .htm       — HTML stripped to text
    .xml              — XML text content
    .toml, .ini, .cfg — config files (plain text)
    .py               — Python (AST-aware: classes + functions as logical units)
    .sql              — SQL (statement-aware: CREATE/INSERT/SELECT blocks)
    .js, .ts, .jsx, .tsx, .go, .rb, .java,
    .cs, .cpp, .c, .rs, .php, .sh, .bash, .zsh,
    .r, .swift, .kt, .scala — code files (plain text)

Dependencies:
    pdfplumber is required for PDF extraction.
    openpyxl is required for .xlsx/.xls extraction.
    python-docx is required for .docx extraction.
"""

from __future__ import annotations

import ast
import re
from pathlib import Path

# Code and config file extensions — read as plain text
_CODE_EXTENSIONS = {
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".sql",
    ".go",
    ".rb",
    ".java",
    ".cs",
    ".cpp",
    ".c",
    ".rs",
    ".php",
    ".sh",
    ".bash",
    ".zsh",
    ".r",
    ".swift",
    ".kt",
    ".scala",
    ".toml",
    ".ini",
    ".cfg",
    ".text",
}


def _guess_source_type(path: Path) -> str:
    """Return a short source type string based on file extension."""
    ext = path.suffix.lower()
    mapping = {
        ".pdf": "pdf",
        ".md": "markdown",
        ".txt": "text",
        ".text": "text",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".json": "json",
        ".csv": "csv",
        ".xlsx": "excel",
        ".xls": "excel",
        ".docx": "docx",
        ".ipynb": "notebook",
        ".html": "html",
        ".htm": "html",
        ".xml": "xml",
    }
    if ext in _CODE_EXTENSIONS:
        return "code"
    return mapping.get(ext, "text")


def _extract_file_text_locally(path: Path) -> str:
    """
    Extract plain text from a local file using the client's own CPU.

    Supports all common document, spreadsheet, notebook, and code file types.

    Args:
        path: Path to the local file.

    Returns:
        Extracted text as a string.

    Raises:
        ImportError: If a required optional dependency is not installed.
    """
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf_text(path)
    elif ext in (".md", ".txt", ".text"):
        return path.read_text(encoding="utf-8", errors="replace")
    elif ext in (".yaml", ".yml"):
        return _extract_yaml_text(path)
    elif ext == ".json":
        return _extract_json_text(path)
    elif ext == ".csv":
        return _extract_csv_text(path)
    elif ext in (".xlsx", ".xls"):
        return _extract_excel_text(path)
    elif ext == ".docx":
        return _extract_docx_text(path)
    elif ext == ".ipynb":
        return _extract_notebook_text(path)
    elif ext in (".html", ".htm"):
        return _extract_html_text(path)
    elif ext == ".xml":
        return _extract_xml_text(path)
    else:
        # Code files, config files, and any other text-based format
        return path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf_text(path: Path) -> str:
    """
    Extract text from a PDF using pdfplumber (text layer only).

    This is a lightweight extraction — it reads the text layer of the PDF
    without OCR.  For scanned/image-only PDFs, the result may be empty or
    sparse.  In that case, consider using the server-side pipeline
    (fetch_mode="server") which includes RapidOCR fallback.

    Args:
        path: Path to the PDF file.

    Returns:
        Extracted text as a string.

    Raises:
        ImportError: If pdfplumber is not installed.
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber is required for local PDF extraction. "
            "Install it with: pip install pdfplumber\n"
            "Or install the full PDF extras: pip install scrapedatshi[pdf]\n"
            "Alternatively, use fetch_mode='server' to have the server parse the PDF."
        )

    import io

    pdf_bytes = path.read_bytes()
    output_blocks: list[str] = []

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # Extract tables first
            tables = page.extract_tables()
            table_bboxes = (
                [t_obj.bbox for t_obj in page.find_tables()] if tables else []
            )

            if table_bboxes:
                remaining = page
                for bbox in table_bboxes:
                    try:
                        remaining = remaining.outside_bbox(bbox)
                    except Exception:
                        pass
                text_content = (
                    remaining.extract_text(x_tolerance=3, y_tolerance=3) or ""
                )
            else:
                text_content = page.extract_text(x_tolerance=3, y_tolerance=3) or ""

            if text_content.strip():
                output_blocks.extend(text_content.splitlines())

            # Add GFM tables
            for table_rows in tables:
                if not table_rows:
                    continue
                cleaned = [
                    [str(cell).strip() if cell is not None else "" for cell in row]
                    for row in table_rows
                ]
                output_blocks.append("")
                output_blocks.append(_table_to_gfm(cleaned))
                output_blocks.append("")

            output_blocks.append("")  # blank line between pages

    text = "\n".join(output_blocks).strip()

    if not text:
        # PDF has no text layer — warn the user
        import warnings

        warnings.warn(
            f"No text could be extracted from '{path.name}'. "
            "This PDF may be a scanned image-only document. "
            "For OCR support, use fetch_mode='server' which includes RapidOCR fallback.",
            stacklevel=4,
        )

    return text


def _table_to_gfm(rows: list[list[str]]) -> str:
    """Convert a list of rows to a GFM markdown table."""
    if not rows:
        return ""

    col_count = max(len(row) for row in rows)
    padded = [row + [""] * (col_count - len(row)) for row in rows]
    col_widths = [
        max(len(str(padded[r][c])) for r in range(len(padded)))
        for c in range(col_count)
    ]
    col_widths = [max(w, 3) for w in col_widths]

    def fmt_row(row: list[str]) -> str:
        cells = [str(row[c]).ljust(col_widths[c]) for c in range(col_count)]
        return "| " + " | ".join(cells) + " |"

    separator = "| " + " | ".join("-" * w for w in col_widths) + " |"
    result_lines = [fmt_row(padded[0]), separator]
    for row in padded[1:]:
        result_lines.append(fmt_row(row))
    return "\n".join(result_lines)


def _extract_yaml_text(path: Path) -> str:
    """Extract text from a YAML file."""
    try:
        import yaml  # type: ignore

        data = yaml.safe_load(path.read_text(encoding="utf-8", errors="replace"))
        return yaml.dump(data, default_flow_style=False, allow_unicode=True)
    except ImportError:
        # PyYAML not installed — read as plain text
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def _extract_json_text(path: Path) -> str:
    """Extract text from a JSON file."""
    import json

    try:
        data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        return json.dumps(data, indent=2, ensure_ascii=False)
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def _extract_csv_text(path: Path) -> str:
    """Extract text from a CSV file — converts rows to readable text blocks."""
    try:
        import csv
        import io

        text = path.read_text(encoding="utf-8", errors="replace")
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return text
        lines: list[str] = []
        for i, row in enumerate(rows, 1):
            row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            if row_text.strip():
                lines.append(f"Row {i}: {row_text}")
        return "\n".join(lines) if lines else text
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def _extract_excel_text(path: Path) -> str:
    """Extract text from an Excel file (.xlsx/.xls). Requires openpyxl."""
    try:
        import openpyxl  # type: ignore

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        all_text: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_lines: list[str] = [f"## Sheet: {sheet_name}"]
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    sheet_lines.append(row_text)
            if len(sheet_lines) > 1:
                all_text.append("\n".join(sheet_lines))
        wb.close()
        return "\n\n".join(all_text) if all_text else ""
    except ImportError:
        raise ImportError(
            "openpyxl is required for Excel file extraction. "
            "Install it with: pip install openpyxl"
        )
    except Exception as exc:
        return f"[Excel parse error: {exc}]"


def _extract_docx_text(path: Path) -> str:
    """Extract text from a Word document (.docx). Requires python-docx."""
    try:
        import docx  # type: ignore

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs) if paragraphs else ""
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document extraction. "
            "Install it with: pip install python-docx"
        )
    except Exception as exc:
        return f"[DOCX parse error: {exc}]"


def _extract_notebook_text(path: Path) -> str:
    """Extract text from a Jupyter notebook (.ipynb)."""
    try:
        import json as _json

        nb = _json.loads(path.read_text(encoding="utf-8", errors="replace"))
        cells = nb.get("cells", [])
        blocks: list[str] = []
        for cell in cells:
            cell_type = cell.get("cell_type", "")
            source = cell.get("source", [])
            if isinstance(source, list):
                source = "".join(source)
            if not source.strip():
                continue
            if cell_type == "markdown":
                blocks.append(source)
            elif cell_type == "code":
                blocks.append(f"```python\n{source}\n```")
        return "\n\n".join(blocks) if blocks else ""
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def _extract_html_text(path: Path) -> str:
    """Extract plain text from an HTML file (strips tags)."""
    try:
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self._parts: list[str] = []
                self._skip = False

            def handle_starttag(self, tag: str, attrs: list) -> None:
                if tag in ("script", "style"):
                    self._skip = True

            def handle_endtag(self, tag: str) -> None:
                if tag in ("script", "style"):
                    self._skip = False

            def handle_data(self, data: str) -> None:
                if not self._skip and data.strip():
                    self._parts.append(data.strip())

        extractor = _TextExtractor()
        extractor.feed(path.read_text(encoding="utf-8", errors="replace"))
        return " ".join(extractor._parts)
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def _extract_xml_text(path: Path) -> str:
    """Extract text content from an XML file."""
    try:
        from xml.etree import ElementTree as ET

        tree = ET.parse(str(path))
        root = tree.getroot()
        texts = [
            elem.text.strip() for elem in root.iter() if elem.text and elem.text.strip()
        ]
        return "\n".join(texts) if texts else ""
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


# ── AST-aware Python code extraction ─────────────────────────────────────────


def _extract_import_names_sdk(import_lines: list[str]) -> list[str]:
    """Extract the top-level names introduced by import statement lines."""
    names: list[str] = []
    for line in import_lines:
        stripped = line.strip()
        if stripped.startswith("from "):
            after_import = stripped.split(" import ", 1)
            if len(after_import) == 2:
                for part in after_import[1].split(","):
                    part = part.strip()
                    if " as " in part:
                        names.append(part.split(" as ")[-1].strip())
                    else:
                        names.append(part)
        elif stripped.startswith("import "):
            after_import = stripped[len("import ") :].strip()
            for part in after_import.split(","):
                part = part.strip()
                if " as " in part:
                    names.append(part.split(" as ")[-1].strip())
                else:
                    names.append(part)
    return [n for n in names if n]


def _extract_python_ast_units(source: str) -> list[dict]:
    """
    Parse a Python source string with the stdlib ``ast`` module and extract
    top-level classes, functions, and module-level (global scope) code as
    logical units.

    Each unit is a dict with:
        text        — the source code for that unit
        metadata    — {type, name, line_start, line_end, docstring}

    Improvements:
    - Only prepends imports that are actually referenced in each unit (no bloat)
    - Collects all module-level code outside named defs as "global_scope" units
      so no lines are silently dropped
    - Falls back to a single module-level unit if parsing fails
    """
    try:
        tree = ast.parse(source)
    except SyntaxError:
        return [
            {
                "text": source,
                "metadata": {
                    "type": "module",
                    "name": "<module>",
                    "line_start": 1,
                    "line_end": source.count("\n") + 1,
                    "docstring": None,
                },
            }
        ]

    lines = source.splitlines()

    # Collect import entries as (stmt_text, names_introduced) for smart filtering
    import_entries: list[tuple[str, list[str]]] = []
    for node in ast.iter_child_nodes(tree):
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            start = node.lineno - 1
            end = node.end_lineno  # type: ignore[attr-defined]
            stmt_lines = lines[start:end]
            stmt_text = "\n".join(stmt_lines)
            names = _extract_import_names_sdk(stmt_lines)
            import_entries.append((stmt_text, names))

    top_level_defs = [
        node
        for node in ast.iter_child_nodes(tree)
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef))
    ]

    def_ranges: list[tuple[int, int]] = [
        (node.lineno, node.end_lineno)  # type: ignore[attr-defined]
        for node in top_level_defs
    ]

    import_line_set: set[int] = set()
    for node in ast.iter_child_nodes(tree):
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            for ln in range(node.lineno, node.end_lineno + 1):  # type: ignore[attr-defined]
                import_line_set.add(ln)

    units: list[dict] = []

    for node in top_level_defs:
        start_0 = node.lineno - 1
        end_0 = node.end_lineno  # type: ignore[attr-defined]
        unit_source = "\n".join(lines[start_0:end_0])

        docstring: str | None = None
        try:
            docstring = ast.get_docstring(node)
        except Exception:
            pass

        if isinstance(node, ast.ClassDef):
            unit_type = "class"
        elif isinstance(node, ast.AsyncFunctionDef):
            unit_type = "async_function"
        else:
            unit_type = "function"

        relevant_imports: list[str] = []
        for stmt_text, names in import_entries:
            if any(name in unit_source for name in names):
                relevant_imports.append(stmt_text)

        text = (
            "\n".join(relevant_imports) + "\n\n" + unit_source
            if relevant_imports
            else unit_source
        )

        units.append(
            {
                "text": text,
                "metadata": {
                    "type": unit_type,
                    "name": node.name,
                    "line_start": node.lineno,
                    "line_end": end_0,
                    "docstring": docstring,
                },
            }
        )

    # Collect global scope remnants (lines not in any def or import)
    covered_lines: set[int] = set(import_line_set)
    for lo, hi in def_ranges:
        for ln in range(lo, hi + 1):
            covered_lines.add(ln)

    global_lines: list[tuple[int, str]] = []
    for i, line_text in enumerate(lines):
        lineno = i + 1
        if lineno not in covered_lines:
            global_lines.append((lineno, line_text))

    if global_lines:
        blocks: list[list[tuple[int, str]]] = []
        current_block: list[tuple[int, str]] = [global_lines[0]]
        for prev, curr in zip(global_lines, global_lines[1:]):
            if curr[0] == prev[0] + 1:
                current_block.append(curr)
            else:
                blocks.append(current_block)
                current_block = [curr]
        blocks.append(current_block)

        for block in blocks:
            block_text = "\n".join(line for _, line in block).strip()
            if not block_text:
                continue
            units.append(
                {
                    "text": block_text,
                    "metadata": {
                        "type": "global_scope",
                        "name": "<module>",
                        "line_start": block[0][0],
                        "line_end": block[-1][0],
                        "docstring": None,
                    },
                }
            )

    units.sort(key=lambda u: u["metadata"]["line_start"])

    if not units:
        return [
            {
                "text": source,
                "metadata": {
                    "type": "module",
                    "name": "<module>",
                    "line_start": 1,
                    "line_end": len(lines),
                    "docstring": None,
                },
            }
        ]

    return units


# ── Statement-aware SQL extraction ────────────────────────────────────────────

# Regex that matches the start of a new SQL statement block
_SQL_STATEMENT_RE = re.compile(
    r"^\s*("
    r"CREATE\s+(OR\s+REPLACE\s+)?(TABLE|VIEW|INDEX|PROCEDURE|FUNCTION|TRIGGER|SCHEMA|DATABASE|SEQUENCE|TYPE)"
    r"|ALTER\s+(TABLE|VIEW|INDEX|PROCEDURE|FUNCTION|SCHEMA|DATABASE|SEQUENCE|TYPE)"
    r"|DROP\s+(TABLE|VIEW|INDEX|PROCEDURE|FUNCTION|TRIGGER|SCHEMA|DATABASE|SEQUENCE|TYPE)"
    r"|INSERT\s+INTO"
    r"|UPDATE\s+\w"
    r"|DELETE\s+FROM"
    r"|SELECT\b"
    r"|WITH\b"
    r"|MERGE\b"
    r"|TRUNCATE\b"
    r"|GRANT\b"
    r"|REVOKE\b"
    r"|COMMENT\s+ON"
    r")",
    re.IGNORECASE | re.MULTILINE,
)


def _extract_sql_units(source: str) -> list[dict]:
    """
    Split a SQL source string into logical statement blocks.

    Uses regex to detect the start of each major SQL statement (CREATE TABLE,
    INSERT INTO, SELECT, etc.) and groups lines between statement boundaries.

    Each unit is a dict with:
        text        — the SQL text for that statement block
        metadata    — {type: "sql_statement", statement_index, statement_type}

    If no statement boundaries are found, returns a single unit with the full
    source so the caller can fall back gracefully.
    """
    lines = source.splitlines(keepends=True)
    if not lines:
        return [
            {
                "text": source,
                "metadata": {
                    "type": "sql_statement",
                    "statement_index": 0,
                    "statement_type": "unknown",
                },
            }
        ]

    # Find line indices where a new statement starts
    boundary_indices: list[int] = []
    for i, line in enumerate(lines):
        if _SQL_STATEMENT_RE.match(line):
            boundary_indices.append(i)

    if not boundary_indices:
        # No recognisable statement boundaries — return whole file
        return [
            {
                "text": source,
                "metadata": {
                    "type": "sql_statement",
                    "statement_index": 0,
                    "statement_type": "unknown",
                },
            }
        ]

    # Build statement blocks
    units: list[dict] = []
    boundaries = boundary_indices + [len(lines)]  # sentinel

    for idx, (start, end) in enumerate(zip(boundaries, boundaries[1:])):
        block_lines = lines[start:end]
        block_text = "".join(block_lines).strip()
        if not block_text:
            continue

        # Detect statement type from first non-whitespace keyword
        first_line = block_lines[0].strip().upper()
        stmt_type = "unknown"
        for keyword in (
            "CREATE TABLE",
            "CREATE OR REPLACE TABLE",
            "CREATE VIEW",
            "CREATE OR REPLACE VIEW",
            "CREATE PROCEDURE",
            "CREATE OR REPLACE PROCEDURE",
            "CREATE FUNCTION",
            "CREATE OR REPLACE FUNCTION",
            "CREATE TRIGGER",
            "CREATE INDEX",
            "ALTER TABLE",
            "ALTER VIEW",
            "DROP TABLE",
            "DROP VIEW",
            "INSERT INTO",
            "UPDATE",
            "DELETE FROM",
            "SELECT",
            "WITH",
            "MERGE",
            "TRUNCATE",
            "GRANT",
            "REVOKE",
        ):
            if first_line.startswith(keyword):
                stmt_type = keyword.lower().replace(" ", "_")
                break

        units.append(
            {
                "text": block_text,
                "metadata": {
                    "type": "sql_statement",
                    "statement_index": idx,
                    "statement_type": stmt_type,
                },
            }
        )

    return (
        units
        if units
        else [
            {
                "text": source,
                "metadata": {
                    "type": "sql_statement",
                    "statement_index": 0,
                    "statement_type": "unknown",
                },
            }
        ]
    )
