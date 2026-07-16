"""
scrapedatshi.pipeline._ingest_helpers
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
Folder ingestion engine and file text extraction.

Supports the following file types for ingest_folder():
    .md, .txt, .text          — plain text / markdown
    .json                     — JSON (with Scrapy array detection)
    .yaml, .yml               — YAML
    .csv                      — CSV rows → text blocks
    .xlsx, .xls               — Excel sheets → text blocks (requires openpyxl)
    .docx                     — Word documents (requires python-docx)
    .ipynb                    — Jupyter notebooks (code + markdown cells)
    .html, .htm               — HTML (stripped to text)
    .xml                      — XML text content
    .toml, .ini, .cfg         — config files (plain text)
    .py, .js, .ts, .jsx, .tsx — code files (plain text)
    .sql, .go, .rb, .java     — code files (plain text)
    .cs, .cpp, .c, .rs, .php  — code files (plain text)
    .sh, .bash, .zsh          — shell scripts (plain text)
    .r, .swift, .kt, .scala   — code files (plain text)

All processing runs on the CLIENT machine.
"""

from __future__ import annotations

import time
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from scrapedatshi.client import ScrapedatshiClient

from scrapedatshi._file_parser import _guess_source_type
from scrapedatshi.models import IngestFolderResult

# ── Supported extensions ──────────────────────────────────────────────────────

# Default extensions for ingest_folder() — covers all supported types
_INGEST_FOLDER_EXTENSIONS = (
    # Documents
    ".md",
    ".txt",
    ".text",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".xlsx",
    ".xls",
    ".docx",
    ".ipynb",
    ".html",
    ".htm",
    ".xml",
    ".toml",
    ".ini",
    ".cfg",
    # Code files
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".sql",
    ".go",
    ".rb",
    ".java",
    ".cs",
    ".cpp",
    ".c",
    ".rs",
    ".php",
    ".sh",
    ".bash",
    ".zsh",
    ".r",
    ".swift",
    ".kt",
    ".scala",
)

# Keys to look for when extracting text from a JSON item (Scrapy/crawler output)
_JSON_TEXT_KEYS = ("text", "content", "html", "body", "markdown", "description")

# Max backoff sleep in seconds for 429 rate limit handling
_MAX_BACKOFF_SLEEP = 60.0


# ── File text extraction ──────────────────────────────────────────────────────


def _extract_text_from_file(
    path: Path, json_text_keys: tuple[str, ...]
) -> list[tuple[str, str]]:
    """
    Extract (text, source_url) pairs from a single file.

    Handles all supported file types. For JSON files, detects Scrapy/crawler
    array exports and yields one entry per item.

    Returns a list of (text, source_url) tuples.
    """
    import json as _json

    ext = path.suffix.lower()
    source_url = f"file://{path.name}"

    # ── JSON ──────────────────────────────────────────────────────────────────
    if ext == ".json":
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            data = _json.loads(raw)
        except Exception:
            return [(path.read_text(encoding="utf-8", errors="replace"), source_url)]

        if isinstance(data, list):
            results: list[tuple[str, str]] = []
            for item in data:
                if isinstance(item, dict):
                    text = None
                    for key in json_text_keys:
                        val = item.get(key)
                        if val and isinstance(val, str) and val.strip():
                            text = val
                            break
                    if text is None:
                        text = _json.dumps(item, ensure_ascii=False)
                    item_url = item.get("url") or item.get("link") or source_url
                    results.append((text, str(item_url)))
                elif isinstance(item, str) and item.strip():
                    results.append((item, source_url))
            return results if results else [(raw, source_url)]

        elif isinstance(data, dict):
            for key in json_text_keys:
                val = data.get(key)
                if val and isinstance(val, str) and val.strip():
                    item_url = data.get("url") or data.get("link") or source_url
                    return [(val, str(item_url))]
            return [(_json.dumps(data, ensure_ascii=False, indent=2), source_url)]

        elif isinstance(data, str):
            return [(data, source_url)]

        return [(raw, source_url)]

    # ── YAML ──────────────────────────────────────────────────────────────────
    elif ext in (".yaml", ".yml"):
        try:
            import yaml  # type: ignore

            data = yaml.safe_load(path.read_text(encoding="utf-8", errors="replace"))
            if isinstance(data, str):
                return [(data, source_url)]
            return [
                (
                    yaml.dump(data, default_flow_style=False, allow_unicode=True),
                    source_url,
                )
            ]
        except Exception:
            return [(path.read_text(encoding="utf-8", errors="replace"), source_url)]

    # ── CSV ───────────────────────────────────────────────────────────────────
    elif ext == ".csv":
        try:
            import csv
            import io

            text = path.read_text(encoding="utf-8", errors="replace")
            reader = csv.DictReader(io.StringIO(text))
            rows = list(reader)
            if not rows:
                return [(text, source_url)]
            # Convert each row to a readable text block
            lines: list[str] = []
            for i, row in enumerate(rows, 1):
                row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                if row_text.strip():
                    lines.append(f"Row {i}: {row_text}")
            return [("\n".join(lines), source_url)] if lines else [(text, source_url)]
        except Exception:
            return [(path.read_text(encoding="utf-8", errors="replace"), source_url)]

    # ── Excel ─────────────────────────────────────────────────────────────────
    elif ext in (".xlsx", ".xls"):
        try:
            import openpyxl  # type: ignore

            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            all_text: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_lines: list[str] = [f"## Sheet: {sheet_name}"]
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        sheet_lines.append(row_text)
                if len(sheet_lines) > 1:
                    all_text.append("\n".join(sheet_lines))
            wb.close()
            return (
                [("\n\n".join(all_text), source_url)]
                if all_text
                else [("", source_url)]
            )
        except ImportError:
            return [
                (
                    f"[openpyxl not installed — cannot parse {path.name}. "
                    "Install with: pip install openpyxl]",
                    source_url,
                )
            ]
        except Exception as exc:
            return [(f"[Excel parse error: {exc}]", source_url)]

    # ── Word DOCX ─────────────────────────────────────────────────────────────
    elif ext == ".docx":
        try:
            import docx  # type: ignore

            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            return (
                [("\n\n".join(paragraphs), source_url)]
                if paragraphs
                else [("", source_url)]
            )
        except ImportError:
            return [
                (
                    f"[python-docx not installed — cannot parse {path.name}. "
                    "Install with: pip install python-docx]",
                    source_url,
                )
            ]
        except Exception as exc:
            return [(f"[DOCX parse error: {exc}]", source_url)]

    # ── Jupyter Notebook ──────────────────────────────────────────────────────
    elif ext == ".ipynb":
        try:
            import json as _json2

            nb = _json2.loads(path.read_text(encoding="utf-8", errors="replace"))
            cells = nb.get("cells", [])
            blocks: list[str] = []
            for cell in cells:
                cell_type = cell.get("cell_type", "")
                source = cell.get("source", [])
                if isinstance(source, list):
                    source = "".join(source)
                if not source.strip():
                    continue
                if cell_type == "markdown":
                    blocks.append(source)
                elif cell_type == "code":
                    blocks.append(f"```python\n{source}\n```")
            return [("\n\n".join(blocks), source_url)] if blocks else [("", source_url)]
        except Exception:
            return [(path.read_text(encoding="utf-8", errors="replace"), source_url)]

    # ── HTML ──────────────────────────────────────────────────────────────────
    elif ext in (".html", ".htm"):
        try:
            from html.parser import HTMLParser

            class _TextExtractor(HTMLParser):
                def __init__(self) -> None:
                    super().__init__()
                    self._parts: list[str] = []
                    self._skip = False

                def handle_starttag(self, tag: str, attrs: list) -> None:
                    if tag in ("script", "style"):
                        self._skip = True

                def handle_endtag(self, tag: str) -> None:
                    if tag in ("script", "style"):
                        self._skip = False

                def handle_data(self, data: str) -> None:
                    if not self._skip and data.strip():
                        self._parts.append(data.strip())

            extractor = _TextExtractor()
            extractor.feed(path.read_text(encoding="utf-8", errors="replace"))
            text = " ".join(extractor._parts)
            return [(text, source_url)] if text.strip() else [("", source_url)]
        except Exception:
            return [(path.read_text(encoding="utf-8", errors="replace"), source_url)]

    # ── XML ───────────────────────────────────────────────────────────────────
    elif ext == ".xml":
        try:
            from xml.etree import ElementTree as ET

            tree = ET.parse(str(path))
            root = tree.getroot()
            texts = [
                elem.text.strip()
                for elem in root.iter()
                if elem.text and elem.text.strip()
            ]
            return [("\n".join(texts), source_url)] if texts else [("", source_url)]
        except Exception:
            return [(path.read_text(encoding="utf-8", errors="replace"), source_url)]

    # ── All other types (plain text: .md, .txt, code files, config files) ────
    else:
        return [(path.read_text(encoding="utf-8", errors="replace"), source_url)]


# ── Folder ingestion loops ────────────────────────────────────────────────────


def _ingest_folder_locally(
    *,
    client: "ScrapedatshiClient",
    folder_path: Path,
    embedding_provider: str,
    embedding_api_key: str,
    vector_db: str,
    vector_db_config: dict,
    embedding_model: str | None,
    embedding_endpoint: str | None,
    chunk_size: int,
    overlap: int,
    file_extensions: tuple[str, ...],
    recursive: bool,
    max_files: int | None,
    batch_delay: float,
    json_text_keys: tuple[str, ...],
) -> "IngestFolderResult":
    """
    Synchronous folder ingestion loop.

    Iterates files, extracts text (with JSON array detection), chunks via
    /v1/process-text, then embeds + injects via /v1/ingest. Includes
    exponential backoff on 429 rate limit errors.
    """
    import json as _json
    from scrapedatshi.exceptions import RateLimitError

    files_processed = 0
    files_failed = 0
    total_chunks = 0
    vectors_upserted = 0
    total_credits_used = 0.0
    last_credits_remaining = 0.0
    errors: list[dict] = []

    if recursive:
        all_files = [
            p
            for p in folder_path.rglob("*")
            if p.is_file() and p.suffix.lower() in file_extensions
        ]
    else:
        all_files = [
            p
            for p in folder_path.iterdir()
            if p.is_file() and p.suffix.lower() in file_extensions
        ]

    all_files.sort()
    if max_files is not None:
        all_files = all_files[:max_files]

    embedding_cfg = {"provider": embedding_provider, "api_key": embedding_api_key}
    if embedding_model:
        embedding_cfg["model"] = embedding_model
    if embedding_endpoint:
        embedding_cfg["endpoint"] = embedding_endpoint
    vdb_cfg = {"provider": vector_db, **vector_db_config}

    for file_path in all_files:
        try:
            text_entries = _extract_text_from_file(file_path, json_text_keys)
        except Exception as exc:
            files_failed += 1
            errors.append(
                {"file": str(file_path), "error": f"Text extraction failed: {exc}"}
            )
            continue

        for text, source_url in text_entries:
            if not text or not text.strip():
                continue

            chunk_payload: dict = {
                "url": source_url,
                "text": text,
                "source_type": _guess_source_type(file_path),
            }
            if chunk_size != 512:
                chunk_payload["chunk_size"] = chunk_size
            if overlap != 50:
                chunk_payload["overlap"] = overlap

            backoff = 2.0
            for attempt in range(5):
                try:
                    chunk_data = client._post("/v1/process-text", json=chunk_payload)
                    break
                except RateLimitError:
                    if attempt == 4:
                        raise
                    time.sleep(min(backoff, _MAX_BACKOFF_SLEEP))
                    backoff *= 2
            else:
                files_failed += 1
                errors.append(
                    {
                        "file": str(file_path),
                        "error": "Rate limit exceeded after retries",
                    }
                )
                continue

            chunks = chunk_data.get("chunks", [])
            if not chunks:
                continue

            form_data: dict = {
                "embedding_config": _json.dumps(embedding_cfg),
                "vector_db_config": _json.dumps(vdb_cfg),
            }
            if chunk_size != 512:
                form_data["chunk_size"] = str(chunk_size)
            if overlap != 50:
                form_data["overlap"] = str(overlap)

            text_bytes = text.encode("utf-8")
            mime = (
                "text/markdown" if file_path.suffix.lower() == ".md" else "text/plain"
            )

            backoff = 2.0
            for attempt in range(5):
                try:
                    ingest_data = client._post(
                        "/v1/ingest",
                        files={"files": (file_path.name, text_bytes, mime)},
                        data=form_data,
                    )
                    break
                except RateLimitError:
                    if attempt == 4:
                        raise
                    time.sleep(min(backoff, _MAX_BACKOFF_SLEEP))
                    backoff *= 2
            else:
                files_failed += 1
                errors.append(
                    {
                        "file": str(file_path),
                        "error": "Rate limit exceeded after retries",
                    }
                )
                continue

            total_chunks += ingest_data.get("total_chunks_created", len(chunks))
            vectors_upserted += ingest_data.get("total_vectors_upserted", 0)
            total_credits_used += float(ingest_data.get("credits_used", 0.0))
            last_credits_remaining = float(ingest_data.get("credits_remaining", 0.0))

        files_processed += 1
        if batch_delay > 0:
            time.sleep(batch_delay)

    return IngestFolderResult(
        files_processed=files_processed,
        files_failed=files_failed,
        total_chunks=total_chunks,
        vectors_upserted=vectors_upserted,
        embedding_provider=embedding_provider,
        vector_db_provider=vector_db,
        credits_used=total_credits_used,
        credits_remaining=last_credits_remaining,
        errors=errors,
    )


async def _ingest_folder_locally_async(
    *,
    client: "ScrapedatshiClient",
    folder_path: Path,
    embedding_provider: str,
    embedding_api_key: str,
    vector_db: str,
    vector_db_config: dict,
    embedding_model: str | None,
    embedding_endpoint: str | None,
    chunk_size: int,
    overlap: int,
    file_extensions: tuple[str, ...],
    recursive: bool,
    max_files: int | None,
    batch_delay: float,
    json_text_keys: tuple[str, ...],
) -> "IngestFolderResult":
    """Async version of :func:`_ingest_folder_locally`."""
    import asyncio as _asyncio
    import json as _json
    from scrapedatshi.exceptions import RateLimitError

    files_processed = 0
    files_failed = 0
    total_chunks = 0
    vectors_upserted = 0
    total_credits_used = 0.0
    last_credits_remaining = 0.0
    errors: list[dict] = []

    if recursive:
        all_files = [
            p
            for p in folder_path.rglob("*")
            if p.is_file() and p.suffix.lower() in file_extensions
        ]
    else:
        all_files = [
            p
            for p in folder_path.iterdir()
            if p.is_file() and p.suffix.lower() in file_extensions
        ]

    all_files.sort()
    if max_files is not None:
        all_files = all_files[:max_files]

    embedding_cfg = {"provider": embedding_provider, "api_key": embedding_api_key}
    if embedding_model:
        embedding_cfg["model"] = embedding_model
    if embedding_endpoint:
        embedding_cfg["endpoint"] = embedding_endpoint
    vdb_cfg = {"provider": vector_db, **vector_db_config}

    for file_path in all_files:
        try:
            text_entries = await _asyncio.to_thread(
                _extract_text_from_file, file_path, json_text_keys
            )
        except Exception as exc:
            files_failed += 1
            errors.append(
                {"file": str(file_path), "error": f"Text extraction failed: {exc}"}
            )
            continue

        for text, source_url in text_entries:
            if not text or not text.strip():
                continue

            chunk_payload: dict = {
                "url": source_url,
                "text": text,
                "source_type": _guess_source_type(file_path),
            }
            if chunk_size != 512:
                chunk_payload["chunk_size"] = chunk_size
            if overlap != 50:
                chunk_payload["overlap"] = overlap

            backoff = 2.0
            for attempt in range(5):
                try:
                    chunk_data = await client._post_async(
                        "/v1/process-text", json=chunk_payload
                    )
                    break
                except RateLimitError:
                    if attempt == 4:
                        raise
                    await _asyncio.sleep(min(backoff, _MAX_BACKOFF_SLEEP))
                    backoff *= 2
            else:
                files_failed += 1
                errors.append(
                    {
                        "file": str(file_path),
                        "error": "Rate limit exceeded after retries",
                    }
                )
                continue

            chunks = chunk_data.get("chunks", [])
            if not chunks:
                continue

            form_data: dict = {
                "embedding_config": _json.dumps(embedding_cfg),
                "vector_db_config": _json.dumps(vdb_cfg),
            }
            if chunk_size != 512:
                form_data["chunk_size"] = str(chunk_size)
            if overlap != 50:
                form_data["overlap"] = str(overlap)

            text_bytes = text.encode("utf-8")
            mime = (
                "text/markdown" if file_path.suffix.lower() == ".md" else "text/plain"
            )

            backoff = 2.0
            for attempt in range(5):
                try:
                    ingest_data = await client._post_async(
                        "/v1/ingest",
                        files={"files": (file_path.name, text_bytes, mime)},
                        data=form_data,
                    )
                    break
                except RateLimitError:
                    if attempt == 4:
                        raise
                    await _asyncio.sleep(min(backoff, _MAX_BACKOFF_SLEEP))
                    backoff *= 2
            else:
                files_failed += 1
                errors.append(
                    {
                        "file": str(file_path),
                        "error": "Rate limit exceeded after retries",
                    }
                )
                continue

            total_chunks += ingest_data.get("total_chunks_created", len(chunks))
            vectors_upserted += ingest_data.get("total_vectors_upserted", 0)
            total_credits_used += float(ingest_data.get("credits_used", 0.0))
            last_credits_remaining = float(ingest_data.get("credits_remaining", 0.0))

        files_processed += 1
        if batch_delay > 0:
            await _asyncio.sleep(batch_delay)

    return IngestFolderResult(
        files_processed=files_processed,
        files_failed=files_failed,
        total_chunks=total_chunks,
        vectors_upserted=vectors_upserted,
        embedding_provider=embedding_provider,
        vector_db_provider=vector_db,
        credits_used=total_credits_used,
        credits_remaining=last_credits_remaining,
        errors=errors,
    )
